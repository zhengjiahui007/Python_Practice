# pip install python-docx

def first():
    from docx import Document

    doc1 = Document()
    doc1.add_heading('如何使用 Python 创建和操作　Word',0)
    doc1.save('./create_data/16_word创建.docx')
def start1():
    from docx import Document

    doc1 = Document()
    # 增加标题
    doc1.add_heading('如何使用 Python 创建和操作　Word',0)
    
    doc1.add_heading('Python 操作 Word 1级标题',level = 1)
    doc1.add_heading('Python 操作 Word 2级标题',level = 2)
    doc1.add_heading('Python 操作 Word 3级标题',level = 3)

    doc1.add_paragraph('这个是段落内容')
    doc1.add_paragraph('这个是段落内容','Title')  # 等于 0级标题
    
    # 增加段落
    paragraph= doc1.add_paragraph('Python是一种面向对象的解释型计算机程序设计语言,由荷兰人Guido van Rossum于1989年发明,Python的定位是“优雅”、“明确”、“简单”,所以Python程序看上去总是...')
    paragraph.add_run('注意：这个是后加入的内容！！！')

    doc1.add_paragraph('Python是一种面向对象的解释型计算机程序设计语言,由荷兰人Guido van Rossum于1989年发明,Python的定位是“优雅”、“明确”、“简单”,所以Python程序看上去总是...')
    
    # 增加列表
    doc1.add_paragraph('哪个不是动物：')
    doc1.add_paragraph('喜洋洋',style ='List Bullet')
    doc1.add_paragraph('懒洋洋',style ='List Bullet')
    doc1.add_paragraph('沸羊羊',style ='List Bullet')
    doc1.add_paragraph('苹果',style ='List Bullet')
    doc1.add_paragraph('灰太狼',style ='List Bullet')

    doc1.add_paragraph('今年的学习计划：')
    doc1.add_paragraph('Python',style ='List Number')
    doc1.add_paragraph('HTML',style ='List Number')
    doc1.add_paragraph('JavaScript',style ='List Number')
    doc1.add_paragraph('Vue',style ='List Number')
    doc1.add_paragraph('Flask',style ='List Number')

    # 增加引用
    doc1.add_paragraph('这个一个引用', style= 'Intense Quote')
    # 增加图片
    from docx.shared import Inches  # 英寸
    pic = doc1.add_picture('./base_data/backg.jpg')
    height = pic.height
    width = pic.width
    p_width = doc1.sections[0].page_width
    sc = (p_width/10-doc1.sections[0].left_margin/10*2)/(width/10)
    pic.width = int(width*sc)
    pic.height = int(height*sc)

    # 增加表格
    table = doc1.add_table(rows=1,cols=3)
    cells = table.rows[0].cells
    cells[0].text = '编号'
    cells[1].text = '姓名'
    cells[2].text = '职业'
    
    data = [
        [1,'吕小布','将军'],
        [2,'诸葛亮','军师'],
        [3,'刘备','主公'],
    ]
    for i,n,w in data:
        cells = table.add_row().cells   # 增加一行表格
        cells[0].text = str(i)  # 不可以写数字
        cells[1].text = n
        cells[2].text = w

    doc1.save('./create_data/16_word填写内容.docx')

def start3():
    from docx.oxml.ns import qn
if __name__ == "__main__":
    # first()
    start1()
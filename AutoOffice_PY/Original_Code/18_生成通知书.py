from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn
def create_doc(car_no,year,month,day,hour,minute,money,type_info):
    doc1 = Document()
    title = doc1.add_paragraph()
    p = title.add_run('车辆违章处罚通知书')
    p.font.size = Pt(30)
    p.font.color.rgb = RGBColor(255,0,0)
    p.font.name=''
    p._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info =f'''京A{car_no}车于{year}年{month}月{day}日{hour}时{minute}分在营运过程中出现 {type_info}（违章）现象，公司按照安全法规和公司相关制度规定决定对该车驾驶员处以{money}元罚款，要求你在今后的营运过程中严格按照相关法律法规运行。（注：罚款金额请在返程后立即到公司缴纳）  
                                     驾驶员签字：                                         年   月   日  

    '''  
    content = doc1.add_paragraph()
    p2 = content.add_run(info)
    content.paragraph_format.first_line_indent = Inches(0.25)
    doc1.save('./create_data/17_通知书.docx')

if __name__ == "__main__":
    car_no ='123456'
    year = 2030
    month = 8
    day = 8
    hour = 12
    minute = 0
    money = 200
    type_info='违反禁令'
    # [[],[],[]]
    create_doc(car_no,year,month,day,hour,minute,money,type_info)
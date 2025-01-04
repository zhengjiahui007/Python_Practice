


def use_style():
    from docx import Document
    from docx.shared import Inches, Pt,RGBColor
    from docx.oxml.ns import qn
    doc1 = Document()


    p1 = doc1.add_paragraph('这是段落1：\n')
    run = p1.add_run('这是内容1.1\n')
    run.bold = True
    p1.add_run('这是内容1.2\n').italic = True
    p1.add_run('这是内容1.3\n').font.size = Pt(26)
    p1.add_run('这是内容1.4\n').font.strike = True
    p1.add_run('这是内容1.5\n').font.shadow = True
    p1.add_run('这是内容1.6\n').font.color.rgb = RGBColor(255,130,71)
    p1.add_run('Test font style 1.7\n').font.name='微软雅黑'
    run = p1.add_run('测试字体 1.8\n')
    run.font.name =''
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc1.add_paragraph('这是段落2:').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc1.add_paragraph('Python是一种面向对象的解释型计算机程序设计语言,由荷兰人Guido van Rossum于1989年发明,Python的定位是“优雅”、“明确”、“简单”,所以Python程序看上去总是...').paragraph_format.left_indent=Inches(0.5)
    doc1.add_paragraph('Python是一种面向对象的解释型计算机程序设计语言,由荷兰人Guido van Rossum于1989年发明,Python的定位是“优雅”、“明确”、“简单”,所以Python程序看上去总是...').paragraph_format.first_line_indent =Inches(0.5)
    
    doc1.add_paragraph('这是段落3:').paragraph_format.space_before = Pt(20)
    doc1.add_paragraph('这是段落4:').paragraph_format.space_after= Pt(20)
    doc1.add_paragraph('这是段落5:')
   
    doc1.add_paragraph('Python是一种面向对象的解释型计算机程序设计语言,由荷兰人Guido van Rossum于1989年发明,Python的定位是“优雅”、“明确”、“简单”,所以Python程序看上去总是...').paragraph_format.line_spacing = 3
    

    doc1.save('./create_data/17_设置样式.docx')
if __name__ == "__main__":
    use_style()